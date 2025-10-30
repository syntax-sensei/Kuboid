import os
import io
import sys
import logging
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import asyncio
from datetime import datetime, timezone, timedelta
from urllib.parse import urlparse
from uuid import uuid4

# Document processing
import PyPDF2
from docx import Document
import openpyxl
from pptx import Presentation
import pandas as pd
import trafilatura

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LangChainDocument

# Qdrant imports
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue

# Supabase imports
from supabase import create_client, Client

# FastAPI imports
from fastapi import FastAPI, HTTPException, BackgroundTasks, Header
import hmac
import hashlib

BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.append(str(BASE_DIR))

try:
    from analytics import router as analytics_router
except ImportError as analytics_import_error:
    logging.warning(f"Failed to import analytics router: {analytics_import_error}")
    analytics_router = None

from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel

# Import config (try relative/package/absolute imports to be robust whether run as module or script)
try:
    # Preferred: relative import when running as a package (uvicorn backend.RAG.docs:app)
    from .config import Config
except Exception:
    try:
        # Try package-style import
        from RAG.config import Config
    except Exception:
        # Fallback: add current directory to sys.path and import by module name
        current_dir = Path(__file__).resolve().parent
        if str(current_dir) not in sys.path:
            sys.path.insert(0, str(current_dir))
        from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# Initialize clients using config
supabase: Client = create_client(Config.SUPABASE_URL, Config.SUPABASE_SERVICE_ROLE_KEY)
# Initialize Qdrant client with optional API key (for Qdrant Cloud or secured instances)
try:
    qdrant_client = QdrantClient(url=Config.QDRANT_URL, api_key=getattr(Config, "QDRANT_API_KEY", None))
except Exception:
    # Fallback to basic initialization if something unexpected happens
    qdrant_client = QdrantClient(url=Config.QDRANT_URL)
embeddings = OpenAIEmbeddings(
    openai_api_key=Config.OPENAI_API_KEY, model="text-embedding-3-large"
)

# Collection name for Qdrant
COLLECTION_NAME = Config.COLLECTION_NAME
URL_ACTIVITY_TABLE = "url_ingestion_activity"
PROCESSING_STATUS_TABLE = "document_processing_status"

WIDGET_SCRIPT_BASE_URL = os.getenv("WIDGET_SCRIPT_BASE_URL", "http://localhost:8000")
FRONTEND_WIDGET_SCRIPT = os.getenv("FRONTEND_WIDGET_SCRIPT", "/widget.js")


class DocumentProcessor:
    """Handles document processing pipeline"""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )

    def extract_text_from_file(self, file_content: bytes, file_name: str) -> str:
        """Extract text from various file formats"""
        file_extension = Path(file_name).suffix.lower()

        try:
            if file_extension == ".pdf":
                return self._extract_from_pdf(file_content)
            elif file_extension == ".docx":
                return self._extract_from_docx(file_content)
            elif file_extension == ".xlsx":
                return self._extract_from_xlsx(file_content)
            elif file_extension == ".txt":
                return file_content.decode("utf-8")

            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {str(e)}")
            return ""

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_from_xlsx(self, file_content: bytes) -> str:
        """Extract text from XLSX"""
        workbook = openpyxl.load_workbook(io.BytesIO(file_content))
        text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                text += " ".join(str(cell) for cell in row if cell) + "\n"
        return text

    def _extract_from_url(self, url: str) -> str:
        """Extract text from URL"""
        downloaded = trafilatura.fetch_url(url)
        extracted_text = trafilatura.extract(downloaded)
        return extracted_text

    def chunk_text(
        self, text: str, metadata: Dict[str, Any]
    ) -> List[LangChainDocument]:
        """Split text into chunks"""
        if not text.strip():
            return []

        # Create LangChain document
        doc = LangChainDocument(page_content=text, metadata=metadata)

        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        return chunks

    async def generate_embeddings(
        self, chunks: List[LangChainDocument]
    ) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        if not chunks:
            return []

        texts = [chunk.page_content for chunk in chunks]
        embeddings_list = await embeddings.aembed_documents(texts)
        return embeddings_list

    async def store_in_qdrant(
        self,
        chunks: List[LangChainDocument],
        embeddings_list: List[List[float]],
        document_id: str,
    ):
        """Store chunks and embeddings in Qdrant"""
        if not chunks or not embeddings_list:
            return

        # Ensure collection exists
        try:
            qdrant_client.get_collection(COLLECTION_NAME)
        except:
            qdrant_client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=1536, distance=Distance.COSINE),
            )
            # Ensure payload indexes exist for owner/widget fields so filtered search works.
            # Qdrant requires indexes for keyword filtering. Create them if possible; ignore failures.
            try:
                qdrant_client.create_payload_index(collection_name=COLLECTION_NAME, field_name="owner_id", field_schema="keyword", wait=True)
            except Exception:
                try:
                    qdrant_client.create_payload_index(collection_name=COLLECTION_NAME, field_name="owner_id", field_type="keyword", wait=True)
                except Exception:
                    logger.info("Could not create payload index for 'owner_id'")

            try:
                qdrant_client.create_payload_index(collection_name=COLLECTION_NAME, field_name="widget_id", field_schema="keyword", wait=True)
            except Exception:
                try:
                    qdrant_client.create_payload_index(collection_name=COLLECTION_NAME, field_name="widget_id", field_type="keyword", wait=True)
                except Exception:
                    logger.info("Could not create payload index for 'widget_id'")

        # Generate valid UUIDs for point IDs
        import uuid

        # Prepare points for Qdrant
        points = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings_list)):
            # Generate a valid UUID for each point
            point_id = str(uuid.uuid4())
            # Extract canonical owner information from chunk metadata when available
            owner_id = None
            try:
                meta = chunk.metadata or {}
                # Extract owner information from metadata
                owner_id = meta.get("user_id")
            except Exception:
                meta = chunk.metadata or {}

            payload = {
                "text": chunk.page_content,
                "metadata": meta,
                "document_id": document_id,
                "chunk_index": i,
                "created_at": datetime.now().isoformat(),
            }

            # Add top-level owner_id for reliable filtering
            if owner_id:
                payload["owner_id"] = owner_id
            # If this chunk was produced in a widget-scoped ingestion, preserve widget_id
            try:
                widget_id_meta = meta.get("widget_id")
                if widget_id_meta:
                    payload["widget_id"] = widget_id_meta
            except Exception:
                pass

            point = PointStruct(id=point_id, vector=embedding, payload=payload)
            points.append(point)

        # Insert points into Qdrant
        try:
            # Log a sample of payload keys for diagnostics
            try:
                sample_payloads = [p.payload for p in points[:3]]
                logger.info("Upserting %d points. Sample payload keys: %s", len(points), [list((pp or {}).keys()) for pp in sample_payloads])
            except Exception:
                pass

            qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)
            logger.info(f"Stored {len(points)} chunks for document {document_id}")
        except Exception as exc:
            logger.exception("Failed to upsert points to Qdrant: %s", exc)


class IngestionPipeline:
    """Main ingestion pipeline"""

    def __init__(self):
        self.processor = DocumentProcessor()

    async def _record_url_activity(self, payload: Dict[str, Any]) -> None:
        logger.info("Persisting URL activity %s", payload.get("id"))
        if not payload.get("url"):
            raise ValueError("URL is required for activity records")
        try:
            supabase.table(URL_ACTIVITY_TABLE).upsert(
                payload, on_conflict="id"
            ).execute()
        except Exception as exc:
            logger.error("Failed to persist URL activity: %s", exc)
            raise

    async def _record_url_activity_start(
        self,
        request_id: str,
        url: str,
        user_id: str | None = None,
        metadata: Dict[str, Any] | None = None,
    ):
        payload = {
            "id": request_id,
            "url": url,
            "status": "processing",
            "started_at": datetime.now(timezone.utc).isoformat(),
        }

        if user_id is not None:
            payload["user_id"] = user_id

        if metadata is not None:
            payload["metadata"] = metadata
        await self._record_url_activity(payload)

    async def _record_url_activity_result(
        self,
        request_id: str,
        status: str,
        *,
        url: str,
        chunks_created: int | None = None,
        error: str | None = None,
        metadata: Dict[str, Any] | None = None,
        user_id: str | None = None,
    ):
        payload = {
            "id": request_id,
            "url": url,
            "status": status,
            "completed_at": datetime.now(timezone.utc).isoformat(),
        }
        if chunks_created is not None:
            payload["chunks_created"] = chunks_created

        if error is not None:
            payload["error"] = error[:500]

        if metadata is not None:
            payload["metadata"] = metadata

        if user_id is not None:
            payload["user_id"] = user_id

        await self._record_url_activity(payload)

    async def process_document_from_url(
        self, url: str, user_id: str | None = None
    ) -> Dict[str, Any]:
        """Process content fetched directly from a URL"""
        try:
            logger.info(f"🌐 Processing URL: {url}")

            text = self.processor._extract_from_url(url)
            if not text or not text.strip():
                raise Exception("No text extracted from URL")

            metadata = {
                "source": "url",
                "url": url,
                "processed_at": datetime.now().isoformat(),
                "text_length": len(text),
            }

            if user_id is not None:
                metadata["user_id"] = user_id

            chunks = self.processor.chunk_text(text, metadata)
            logger.info(f"📦 Created {len(chunks)} chunks from URL")

            for chunk in chunks:
                chunk.metadata["chunk_count"] = len(chunks)

            embeddings_list = await self.processor.generate_embeddings(chunks)
            logger.info(f"✅ Generated {len(embeddings_list)} embeddings")

            parsed_url = urlparse(url)
            path_part = parsed_url.path.replace("/", "_") or "root"
            document_id = f"url_{parsed_url.netloc}_{path_part}"

            await self.processor.store_in_qdrant(chunks, embeddings_list, document_id)

            logger.info(f"🎉 Successfully processed URL: {url}")
            return {
                "status": "success",
                "document_id": document_id,
                "chunks_created": len(chunks),
                "text_length": len(text),
                "url": url,
            }

        except Exception as e:
            logger.error(f"❌ Error processing URL {url}: {str(e)}")
            return {"status": "error", "error": str(e)}

    async def process_document_by_path(self, file_path: str) -> Dict[str, Any]:
        """Process a document by its exact path in Supabase storage"""
        try:
            logger.info(f"🚀 Processing document: {file_path}")

            # Mark processing as started in status table (best-effort)
            try:
                document_id = file_path.replace("/", "_").replace("-", "_")
                self._update_processing_status(file_path=document_id, status="processing")
            except Exception:
                # Non-fatal: continue processing even if status update fails
                logger.debug("Failed to record processing start; continuing")

            # Download file from Supabase
            response = supabase.storage.from_("Docs").download(file_path)
            if not response:
                raise Exception(f"Failed to download file: {file_path}")

            # Extract filename from path
            file_name = file_path.split("/")[-1] if "/" in file_path else file_path
            logger.info(f"📁 Processing file: {file_name}")

            # Extract text
            text = self.processor.extract_text_from_file(response, file_name)
            if not text.strip():
                raise Exception("No text extracted from document")

            logger.info(f"📄 Extracted {len(text)} characters")

            # Create metadata
            metadata = {
                "source": "file",
                "file_name": file_name,
                "file_path": file_path,
                "processed_at": datetime.now().isoformat(),
                "text_length": len(text),
            }
            # Attempt to derive user_id from the file path prefix if present (e.g. "<user_id>/filename.pdf")
            try:
                if "/" in file_path:
                    possible_user = file_path.split("/", 1)[0]
                    # Basic sanity: avoid empty and placeholder names
                    if possible_user and not possible_user.startswith("."):
                        metadata["user_id"] = possible_user
            except Exception:
                # If anything goes wrong, don't block processing — leave metadata as-is
                pass

            # Chunk text
            chunks = self.processor.chunk_text(text, metadata)
            logger.info(f"📦 Created {len(chunks)} chunks")

            for chunk in chunks:
                chunk.metadata["chunk_count"] = len(chunks)

            # Generate embeddings
            logger.info(f"🧠 Generating embeddings...")
            embeddings_list = await self.processor.generate_embeddings(chunks)
            logger.info(f"✅ Generated {len(embeddings_list)} embeddings")

            # Store in Qdrant
            document_id = file_path.replace("/", "_").replace("-", "_")
            logger.info(f"💾 Storing in Qdrant...")
            await self.processor.store_in_qdrant(chunks, embeddings_list, document_id)

            logger.info(f"🎉 Successfully processed {file_name}")
            # Mark processing as completed
            try:
                self._update_processing_status(file_path=document_id, status="processed", chunks_count=len(chunks))
            except Exception:
                logger.warning("Failed to persist processing completion status")
            return {
                "status": "success",
                "document_id": document_id,
                "chunks_created": len(chunks),
                "text_length": len(text),
                "file_name": file_name,
            }

        except Exception as e:
            logger.error(f"❌ Error processing {file_path}: {str(e)}")
            # Mark processing as errored
            try:
                document_id = file_path.replace("/", "_").replace("-", "_")
                self._update_processing_status(file_path=document_id, status="error", error=str(e))
            except Exception:
                logger.debug("Failed to persist error status")

            return {"status": "error", "error": str(e)}

    async def is_document_processed(self, file_path: str) -> bool:
        """Check if a document has already been processed"""
        try:
            # Prefer authoritative record in Supabase processing status table
            document_id = file_path.replace("/", "_").replace("-", "_")
            try:
                resp = (
                    supabase.table(PROCESSING_STATUS_TABLE)
                    .select("status")
                    .eq("id", document_id)
                    .limit(1)
                    .execute()
                )
                rows = resp.data if resp and hasattr(resp, "data") else []
                if rows and rows[0].get("status") == "processed":
                    return True
            except Exception as exc:
                logger.debug(f"Processing status table check failed: {exc}")

            # Fallback: check Qdrant for any point with matching document_id
            try:
                q_filter = Filter(must=[FieldCondition(key="document_id", match=MatchValue(value=document_id))])
                points, _ = qdrant_client.scroll(collection_name=COLLECTION_NAME, scroll_filter=q_filter, limit=1)
                return len(points) > 0
            except Exception as exc:
                logger.warning(f"Could not check Qdrant for document presence: {exc}")
                return False

        except Exception as e:
            logger.warning(f"Could not check if document is processed: {e}")
            return False

    async def process_all_documents(
        self, force_reprocess: bool = False, user_id: str | None = None
    ) -> Dict[str, Any]:
        """Process documents in the bucket. If user_id is provided, only process files in that user's folder.
        Skips already processed ones when force_reprocess is False.
        """
        try:
            logger.info("🔄 Starting batch processing of all documents")

            # Determine the path to list. If user_id is provided, list inside that user's folder.
            list_path = user_id or ""

            # Get all files from Supabase under list_path
            all_files = supabase.storage.from_("Docs").list(list_path)

            # Supabase may return folder entries (directories) as list items without an 'id'.
            # Also Supabase may include a placeholder file named '.emptyFolderPlaceholder' when a folder
            # is empty — skip those as they are not real documents.
            file_items = [
                f
                for f in (all_files or [])
                if f.get("id") and not (f.get("name") or "").endswith(".emptyFolderPlaceholder")
            ]
            logger.info(
                f"📋 Found {len(all_files or [])} items at '{list_path}', {len(file_items)} files to process (placeholders skipped)"
            )

            results = []
            skipped = 0

            for file_info in file_items:
                # file_info['name'] may be returned as a basename when listing a folder.
                # Ensure we construct the full path relative to the bucket. If we listed a user folder
                # (list_path != ""), prefix the filename with the folder name when needed.
                raw_name = file_info.get("name") or ""
                if list_path:
                    if raw_name.startswith(list_path + "/") or raw_name == list_path:
                        file_path = raw_name
                    else:
                        file_path = f"{list_path.rstrip('/')}/{raw_name.lstrip('/')}"
                else:
                    file_path = raw_name

                # Check if already processed (unless force reprocess)
                if not force_reprocess and await self.is_document_processed(file_path):
                    logger.info(f"⏭️ Skipping already processed: {file_path}")
                    skipped += 1
                    results.append(
                        {
                            "file_path": file_path,
                            "result": {
                                "status": "skipped",
                                "message": "Already processed",
                            },
                        }
                    )
                    continue

                logger.info(f"🔄 Processing: {file_path}")
                # Process document; points will be tagged with owner/user info only
                result = await self.process_document_by_path(file_path)
                results.append({"file_path": file_path, "result": result})

                # Small delay to avoid overwhelming the system
                await asyncio.sleep(1)

            successful = sum(1 for r in results if r["result"]["status"] == "success")
            failed = sum(1 for r in results if r["result"]["status"] == "error")

            logger.info(
                f"✅ Batch processing complete: {successful} successful, {failed} failed, {skipped} skipped"
            )

            return {
                "status": "completed",
                "total_files": len(results),
                "successful": successful,
                "failed": failed,
                "skipped": skipped,
                "results": results,
            }

        except Exception as e:
            logger.error(f"❌ Batch processing error: {str(e)}")
            return {"status": "error", "error": str(e)}

    def _update_processing_status(
        self, file_path: str, status: str, chunks_count: int = 0, error: str = None
    ):
        """Update processing status in Supabase"""
        try:
            # Upsert a record into Supabase for processing status. This is the
            # authoritative source that we check before reprocessing files.
            payload = {
                "id": file_path,
                "file_path": file_path,
                "status": status,
                "chunks_count": chunks_count,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }

            if status == "processing":
                payload["started_at"] = datetime.now(timezone.utc).isoformat()
            else:
                # For processed or error states mark completed_at
                payload["completed_at"] = datetime.now(timezone.utc).isoformat()

            if error:
                payload["error"] = (error[:500]) if isinstance(error, str) else str(error)

            # Best-effort: if the table doesn't exist this will raise; we don't want
            # this to block ingestion so we log and continue.
            supabase.table(PROCESSING_STATUS_TABLE).upsert(payload, on_conflict="id").execute()
            logger.info(f"Processing status updated for {file_path}: {status}")
        except Exception as e:
            logger.warning(f"Failed to update processing status: {str(e)}")

    async def retrieve(self, query: str, top_k: int = 5, owner_id: str | None = None, widget_id: str | None = None) -> List[Dict[str, Any]]:
        """Retrieve top-k documents for a query.

        Filtering is performed only on the basis of `owner_id` (payload.owner_id).
        If `owner_id` is not provided, no owner/site-level filtering is applied.
        Widget-scoped points (payload.widget_id) are considered only when owner filtering is active
        and widget-specific points exist.
        """
        if not query.strip():
            return []

        query_embedding = await embeddings.aembed_query(query)

        # Decide on metadata filter to restrict results (owner-only filtering)
        qdrant_filter = None
        try:
            if owner_id:
                base_must = [FieldCondition(key="owner_id", match=MatchValue(value=owner_id))]

                # If widget_id provided, check if any points exist with payload.widget_id == widget_id
                if widget_id:
                    try:
                        widget_points, _ = qdrant_client.scroll(
                            collection_name=COLLECTION_NAME,
                            scroll_filter=Filter(must=[FieldCondition(key="widget_id", match=MatchValue(value=widget_id))]),
                            limit=1,
                        )
                        has_widget_scoped = len(widget_points) > 0
                    except Exception:
                        has_widget_scoped = False

                    if has_widget_scoped:
                        # Restrict to widget-scoped points that belong to the owner
                        qdrant_filter = Filter(must=base_must + [FieldCondition(key="widget_id", match=MatchValue(value=widget_id))])
                    else:
                        # No widget-scoped points found; restrict to owner-level points
                        qdrant_filter = Filter(must=base_must)
                else:
                    qdrant_filter = Filter(must=base_must)
            else:
                qdrant_filter = None
        except Exception:
            qdrant_filter = None

        # Execute search with optional filter
        # Log filter for diagnostics
        try:
            logger.info("Qdrant search: collection=%s top_k=%s owner_id=%s widget_id=%s qdrant_filter=%s", COLLECTION_NAME, top_k, owner_id, widget_id, qdrant_filter)
        except Exception:
            pass
        # qdrant-client versions expect the named parameter `query_filter` (not `filter`) for search
        # Use `query_filter` so this works across qdrant-client releases.
        search_results = qdrant_client.search(
            collection_name=COLLECTION_NAME,
            query_vector=query_embedding,
            limit=top_k,
            query_filter=qdrant_filter,
        )
        try:
            logger.info("Qdrant returned %d hits", len(search_results))
        except Exception:
            pass

        documents = []
        for result in search_results:
            payload = result.payload or {}
            documents.append(
                {
                    "text": payload.get("text", ""),
                    "metadata": payload.get("metadata", {}),
                    "score": result.score,
                }
            )
        return documents

    async def generate_answer(
        self,
        query: str,
        context_docs: List[Dict[str, Any]],
        temperature: float = 0.2,
    ) -> str:
        if not context_docs:
            return "I couldn't find relevant information in the knowledge base."

        context_text = "\n\n".join(doc.get("text", "") for doc in context_docs)

        prompt = (
            "You are a helpful assistant that answers questions using the provided context. "
            "If the context does not contain the answer, say that you don't have enough information.\n\n"
            "Context:\n"
            f"{context_text}\n\n"
            "Question:\n"
            f"{query}\n\n"
            "Answer:"
        )

        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=Config.OPENAI_API_KEY)
        completion = await client.responses.create(
            model="gpt-4o-mini",
            input=prompt,
            temperature=temperature,
        )

        answer = completion.output_text.strip()
        return answer


# Initialize pipeline
pipeline = IngestionPipeline()

# FastAPI app
app = FastAPI(title="Document Ingestion Pipeline")

if analytics_router:
    logger.info("Analytics routes enabled")
    app.include_router(analytics_router)
else:
    logger.warning("Analytics routes disabled (module not available)")

# Add CORS middleware
# Validate required config (this will raise on startup if required env vars are missing)
Config.validate()

# Load CORS origins from environment variable `CORS_ORIGINS` (comma-separated, no fallback)
cors_env = Config.CORS_ORIGINS
if not cors_env:
    raise RuntimeError("CORS_ORIGINS environment variable is required and must contain a comma-separated list of origins")

allow_origins = [origin.strip() for origin in cors_env.split(",") if origin.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class IngestionSource(BaseModel):
    type: str
    file_path: str | None = None
    file_name: str | None = None
    url: str | None = None


class ProcessDocumentRequest(BaseModel):
    file_path: str


class UrlIngestionRequest(BaseModel):
    url: str
    request_id: str
    metadata: Dict[str, Any] | None = None


class WidgetTokenRequest(BaseModel):
    widget_id: str
    expires_in: int | None = 3600


class WidgetChatRequest(BaseModel):
    query: str
    history: List[Dict[str, Any]] | None = None
    top_k: int | None = 5
    temperature: float | None = 0.2
    conversation_id: str | None = None


class WidgetCreateRequest(BaseModel):
    site_id: str
    name: str | None = None
    allowed_origins: List[str] | None = None


class WidgetUpdateRequest(BaseModel):
    name: str | None = None
    allowed_origins: List[str] | None = None


from typing import Optional


def _extract_user_id_from_auth(authorization: str | None) -> str:
    try:
        if not authorization or not authorization.startswith("Bearer "):
            return "anonymous"
        token = authorization.split(" ", 1)[1]
        import jwt  # type: ignore

        decoded = jwt.decode(token, options={"verify_signature": False})
        user_id = decoded.get("sub") or decoded.get("user_id") or decoded.get("id")
        return str(user_id) if user_id else "anonymous"
    except Exception as exc:
        logger.warning(f"Failed to extract user id from auth: {exc}")
        return "anonymous"


def _record_chat_turn(
    *,
    owner_id: str,
    conversation_id: str,
    user_message: str,
    assistant_message: str | None,
    status: str,
    metadata: Dict[str, Any] | None = None,
) -> Optional[str]:
    try:
        payload = {
            "conversation_id": conversation_id,
            "owner_id": owner_id,
            "user_message": user_message,
            "assistant_message": assistant_message,
            "status": status,
            "metadata": metadata or {},
        }
        response = supabase.table("chat_turns").insert(payload).execute()
        if response.data:
            return response.data[0].get("id")
        return None
    except Exception as exc:
        logger.warning("Failed to record chat turn: %s", exc, exc_info=True)
        return None


@app.post("/process-document")
async def process_document_endpoint(
    request: ProcessDocumentRequest, background_tasks: BackgroundTasks
):
    """Endpoint to process a document"""
    background_tasks.add_task(
        pipeline.process_document_by_path,
        request.file_path,
    )
    return {"message": "Document processing started"}


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy"}


@app.post("/process-all")
async def process_all_documents(force_reprocess: bool = False):
    """Process all documents in the bucket - Simple and reliable"""
    try:
        result = await pipeline.process_all_documents(force_reprocess=force_reprocess)
        return result
    except Exception as e:
        logger.error(f"Batch processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/process-new-only")
async def process_new_documents(authorization: str = Header(None)):
    """Process only new documents for the authenticated user (skip already processed ones)"""
    try:
        # Require authorization to ensure we only process files for the requesting user
        user_id = _extract_user_id_from_auth(authorization)
        if not authorization or not user_id or user_id == "anonymous":
            raise HTTPException(status_code=401, detail="Missing or invalid authorization token")

        result = await pipeline.process_all_documents(force_reprocess=False, user_id=user_id)
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"New documents processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/process-specific")
async def process_specific_document(request: ProcessDocumentRequest):
    """Process a specific document by its exact path"""
    try:
        result = await pipeline.process_document_by_path(request.file_path)
        return result
    except Exception as e:
        logger.error(f"Specific processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/process-url")
async def process_url_endpoint(
    request: UrlIngestionRequest, authorization: str = Header(None)
):
    """Process content scraped from a URL and record activity"""
    # Extract user_id from Supabase JWT token
    user_id = _extract_user_id_from_auth(authorization)

    await pipeline._record_url_activity_start(
        request.request_id,
        request.url,
        user_id=user_id,
        metadata=request.metadata,
    )
    result = await pipeline.process_document_from_url(request.url, user_id)
    if result.get("status") == "error":
        await pipeline._record_url_activity_result(
            request.request_id,
            "error",
            url=request.url,
            user_id=user_id,
            error=result.get("error"),
            metadata=request.metadata,
        )
        raise HTTPException(status_code=400, detail=result.get("error"))

    await pipeline._record_url_activity_result(
        request.request_id,
        "success",
        url=request.url,
        user_id=user_id,
        chunks_created=result.get("chunks_created"),
        metadata=request.metadata,
    )
    return result


@app.post("/widget/token")
async def create_widget_token(
    request: WidgetTokenRequest,
    origin: str | None = Header(None),
    x_widget_secret: str | None = Header(None),
):
    """Issue a short-lived widget token for a specific widget instance.

    This endpoint requires the browser Origin header to match the widget's
    allowed_origins (stored in the `widgets` table). The token contains both
    `owner_id` and `widget_id` so subsequent requests can be validated.
    """
    try:
        import jwt

        widget_id = request.widget_id
        # Lookup widget record
        try:
            resp = supabase.table("widgets").select("*").eq("id", widget_id).limit(1).execute()
            widget_rows = resp.data if resp and hasattr(resp, "data") else []
            widget = widget_rows[0] if widget_rows else None
        except Exception as exc:
            logger.warning(f"Failed to fetch widget record: {exc}")
            widget = None

        if not widget:
            raise HTTPException(status_code=404, detail="Widget not found")

        allowed_origins = widget.get("allowed_origins") or []
        owner_id = widget.get("owner_id")

        # Authorization modes:
        # 1) Browser: Origin header must be present and match an allowed origin (same as before)
        # 2) Server-to-server: Request may include X-Widget-Secret header matching the widget's secret
        # DB schema may store a hashed secret under `secret_hash` (preferred) or plaintext under `secret` (legacy).
        secret_hash_stored = widget.get("secret_hash")
        legacy_secret_stored = widget.get("secret")

        origin_ok = False
        secret_ok = False

        # Check widget secret if provided (server-to-server flow)
        if x_widget_secret and secret_hash_stored:
            # compute HMAC and compare
            try:
                expected = hmac.new(Config.EMBED_SECRET.encode(), x_widget_secret.encode(), hashlib.sha256).hexdigest()
                if expected == secret_hash_stored:
                    secret_ok = True
            except Exception:
                pass
        elif x_widget_secret and legacy_secret_stored:
            # Backwards compatibility: stored plaintext secret
            if x_widget_secret == legacy_secret_stored:
                secret_ok = True

        # Check origin if provided (browser flow)
        if origin and allowed_origins:
            for allowed in allowed_origins:
                try:
                    if allowed == origin:
                        origin_ok = True
                        break
                except Exception:
                    continue

        if not (origin_ok or secret_ok):
            logger.warning("Widget token request rejected: origin/secret validation failed")
            raise HTTPException(status_code=401, detail="Unauthorized widget token request")

        expires_in = request.expires_in or 60 * 15  # default 15 minutes for widget tokens
        issued_at = datetime.now(timezone.utc)
        expiry = issued_at + timedelta(seconds=expires_in)

        payload = {
            "owner_id": owner_id,
            "widget_id": widget_id,
            "iat": int(issued_at.timestamp()),
            "exp": int(expiry.timestamp()),
        }

        token = jwt.encode(payload, Config.EMBED_SECRET, algorithm="HS256")
        return {"token": token, "expires_at": expiry.isoformat()}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create widget token: {e}")
        raise HTTPException(status_code=500, detail="Could not generate widget token")


@app.post("/widgets")
async def create_widget(request: WidgetCreateRequest, authorization: str = Header(None)):
    """Create a new widget instance for the authenticated owner. Widget id is generated server-side.

    Requires a valid owner JWT in Authorization header (Supabase auth token). Returns the created widget record.
    """
    try:
        owner_id = _extract_user_id_from_auth(authorization)
        if not authorization or owner_id == "anonymous":
            raise HTTPException(status_code=401, detail="Authentication required to create widget")

        import uuid
        widget_id = str(uuid.uuid4())

        payload = {
            "id": widget_id,
            "site_id": request.site_id,
            "owner_id": owner_id,
            "name": request.name,
            "allowed_origins": request.allowed_origins or [],
                # Generate a per-widget secret hash for server-to-server flows (stored in `secret_hash` column)
                "secret_hash": None,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }

        # Ensure the site exists (widgets.site_id has FK to sites.id). If the site
        # doesn't exist, create it for this owner (best-effort) to avoid FK errors.
        try:
            site_resp = supabase.table("sites").select("*").eq("id", request.site_id).limit(1).execute()
            site_rows = site_resp.data if site_resp and hasattr(site_resp, "data") else []
        except Exception as exc:
            logger.warning(f"Failed to check site existence: {exc}")
            site_rows = []

        if not site_rows:
            # Create a minimal site record so widget FK constraint is satisfied
            try:
                site_payload = {
                    "id": request.site_id,
                    "user_id": owner_id,
                    "name": request.site_id,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }
                create_site_resp = supabase.table("sites").insert(site_payload).execute()
                if hasattr(create_site_resp, "error") and create_site_resp.error:
                    logger.error("Failed to create site record: %s", create_site_resp.error)
                    # Let the error surface to client
                    raise Exception(getattr(create_site_resp, "error", "Could not create site"))
                logger.info("Created site record for id=%s", request.site_id)
            except Exception as exc:
                logger.exception("Failed to create site required for widget: %s", exc)
                raise HTTPException(status_code=500, detail=f"Could not create site {request.site_id}: {exc}")

        import secrets
        # create a widget secret and include its HMAC-SHA256 in payload under `secret_hash`
        try:
            secret_val = secrets.token_urlsafe(32)
            # compute HMAC-SHA256 using EMBED_SECRET as key
            secret_hash = hmac.new(
                Config.EMBED_SECRET.encode(), secret_val.encode(), hashlib.sha256
            ).hexdigest()
            payload["secret_hash"] = secret_hash
        except Exception:
            secret_val = None
            secret_hash = None

        try:
            resp = supabase.table("widgets").insert(payload).execute()
            # Log Supabase response for debugging
            logger.info("Supabase insert response: %s", getattr(resp, "data", None))
            if hasattr(resp, "error") and resp.error:
                # Postgrest may return dict-like error
                err = getattr(resp, "error", resp)
                logger.error("Supabase returned error on widget insert: %s", err)
                raise Exception(getattr(resp, "error", "Unknown supabase error"))

            created = resp.data[0] if resp and hasattr(resp, "data") and resp.data else payload
            # Return the secret only at creation time to the owner (do not leak in list endpoints)
            if secret_val:
                created = dict(created)
                created["secret"] = secret_val
                created["secret_persisted"] = True if secret_hash else False
            return {"widget": created}
        except Exception as exc:
            # Handle case where the widgets table schema doesn't include `secret`
            msg = str(exc)
            logger.exception("Failed to insert widget into Supabase: %s", exc)
            if "secret_hash" in msg or "Could not find the 'secret_hash'" in msg or "secret" in msg:
                # Retry insert without `secret_hash` to avoid schema problems, but still return the generated secret
                try:
                    payload_no_secret = dict(payload)
                    payload_no_secret.pop("secret_hash", None)
                    resp2 = supabase.table("widgets").insert(payload_no_secret).execute()
                    if hasattr(resp2, "error") and resp2.error:
                        logger.error("Retry insert without secret_hash also failed: %s", resp2.error)
                        raise Exception(getattr(resp2, "error", "Supabase insert failed"))
                    created = resp2.data[0] if resp2 and hasattr(resp2, "data") and resp2.data else payload_no_secret
                    # Return the generated secret to the caller but note it was not persisted
                    created = dict(created)
                    if secret_val:
                        created["secret"] = secret_val
                        created["secret_persisted"] = False
                    logger.warning("secret_hash column missing in DB; inserted widget without secret_hash. Please run migration to add the column: ALTER TABLE widgets ADD COLUMN secret_hash text;")
                    return {"widget": created}
                except Exception as exc2:
                    logger.exception("Retry insert without secret failed: %s", exc2)
                    raise HTTPException(status_code=500, detail=str(exc2))
            # Other errors: surface to client
            raise HTTPException(status_code=500, detail=msg)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create widget: {e}")
        raise HTTPException(status_code=500, detail="Could not create widget")


@app.get("/widgets")
async def list_widgets(authorization: str = Header(None)):
    """List widgets owned by the authenticated user."""
    try:
        owner_id = _extract_user_id_from_auth(authorization)
        if not authorization or owner_id == "anonymous":
            raise HTTPException(status_code=401, detail="Authentication required")

        resp = supabase.table("widgets").select("*").eq("owner_id", owner_id).execute()
        rows = resp.data if resp and hasattr(resp, "data") else []
        # Do not return secret in list responses for security
        sanitized = []
        for r in rows:
            if isinstance(r, dict):
                r = dict(r)
                # remove any secret material
                r.pop("secret", None)
                r.pop("secret_hash", None)
            sanitized.append(r)
        return {"widgets": sanitized}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to list widgets: {e}")
        raise HTTPException(status_code=500, detail="Could not list widgets")


@app.patch("/widgets/{widget_id}")
async def update_widget(widget_id: str, request: WidgetUpdateRequest, authorization: str = Header(None)):
    """Update widget fields (name, allowed_origins). Owner-only."""
    try:
        owner_id = _extract_user_id_from_auth(authorization)
        if not authorization or owner_id == "anonymous":
            raise HTTPException(status_code=401, detail="Authentication required")

        # Verify ownership
        resp = supabase.table("widgets").select("*").eq("id", widget_id).limit(1).execute()
        rows = resp.data if resp and hasattr(resp, "data") else []
        if not rows:
            raise HTTPException(status_code=404, detail="Widget not found")
        widget = rows[0]
        if str(widget.get("owner_id")) != str(owner_id):
            raise HTTPException(status_code=403, detail="Not authorized to modify this widget")

        update_payload = {}
        if request.name is not None:
            update_payload["name"] = request.name
        if request.allowed_origins is not None:
            update_payload["allowed_origins"] = request.allowed_origins

        if not update_payload:
            return {"widget": widget}

        upd = supabase.table("widgets").update(update_payload).eq("id", widget_id).execute()
        if hasattr(upd, "error") and upd.error:
            logger.error("Supabase error on widget update: %s", upd.error)
            raise Exception(getattr(upd, "error", "Supabase update error"))

        updated = upd.data[0] if upd and hasattr(upd, "data") and upd.data else None
        return {"widget": updated}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update widget: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/widgets/{widget_id}")
async def delete_widget(widget_id: str, authorization: str = Header(None)):
    """Delete a widget. Owner-only."""
    try:
        owner_id = _extract_user_id_from_auth(authorization)
        if not authorization or owner_id == "anonymous":
            raise HTTPException(status_code=401, detail="Authentication required")

        # Verify ownership
        resp = supabase.table("widgets").select("owner_id").eq("id", widget_id).limit(1).execute()
        rows = resp.data if resp and hasattr(resp, "data") else []
        if not rows:
            raise HTTPException(status_code=404, detail="Widget not found")
        if str(rows[0].get("owner_id")) != str(owner_id):
            raise HTTPException(status_code=403, detail="Not authorized to delete this widget")

        res = supabase.table("widgets").delete().eq("id", widget_id).execute()
        if hasattr(res, "error") and res.error:
            logger.error("Supabase error on widget delete: %s", res.error)
            raise Exception(getattr(res, "error", "Supabase delete error"))

        return {"status": "deleted", "widget_id": widget_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to delete widget: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/widget/chat")
async def widget_chat(request: WidgetChatRequest, authorization: str = Header(None)):
    try:
        import jwt

        if not authorization or not authorization.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Missing authorization token")

        token = authorization.split(" ", 1)[1]
        try:
            decoded = jwt.decode(token, Config.EMBED_SECRET, algorithms=["HS256"])
        except jwt.PyJWTError as exc:  # type: ignore
            logger.error(f"Invalid widget token: {exc}")
            raise HTTPException(status_code=401, detail="Invalid authorization token")

        widget_id = decoded.get("widget_id")
        owner_id = decoded.get("owner_id")
        if not widget_id:
            raise HTTPException(status_code=400, detail="Token missing widget_id")

        # Validate widget record and ensure the token's site_id matches the widget
        try:
            resp = supabase.table("widgets").select("*").eq("id", widget_id).limit(1).execute()
            rows = resp.data if resp and hasattr(resp, "data") else []
            widget = rows[0] if rows else None
        except Exception as exc:
            logger.warning(f"Failed to fetch widget for validation: {exc}")
            widget = None

        if not widget:
            logger.error(f"Widget validation failed for widget_id={widget_id}")
            raise HTTPException(status_code=401, detail="Invalid widget token")

        # Prefer owner_id from token; fall back to DB value when token doesn't include it
        if not owner_id:
            owner_id = widget.get("owner_id")

        # Ensure the token (or fallback) matches the widget record
        if str(widget.get("owner_id")) != str(owner_id):
            logger.error(f"Widget owner validation failed for widget_id={widget_id} (expected owner_id={widget.get('owner_id')} got {owner_id})")
            raise HTTPException(status_code=401, detail="Invalid widget token")

        conversation_id = request.conversation_id or decoded.get("conversation_id")
        new_conversation = False
        if not conversation_id:
            conversation_id = str(uuid4())
            new_conversation = True

        started_at = datetime.now(timezone.utc)

        # Restrict retrieval to the widget's owner_id (prefer owner-based filtering) and be widget-aware
        documents = await pipeline.retrieve(request.query, top_k=request.top_k or 5, owner_id=owner_id, widget_id=widget_id)
        answer = await pipeline.generate_answer(
            request.query,
            documents,
            temperature=request.temperature or 0.2,
        )

        elapsed_ms = int(
            (datetime.now(timezone.utc) - started_at).total_seconds() * 1000
        )

        document_refs = [
            {
                "metadata": doc.get("metadata", {}),
                "score": doc.get("score"),
            }
            for doc in documents
        ]

        status = "resolved" if answer else "gap"

        turn_id = None
        if conversation_id:
            turn_id = _record_chat_turn(
                owner_id=owner_id,
                conversation_id=conversation_id,
                user_message=request.query,
                assistant_message=answer,
                status=status,
                metadata={
                    "documents": document_refs,
                    "latency_ms": elapsed_ms,
                    "model": "gpt-4o-mini",
                },
            )

        return {
            "answer": answer,
            "documents": documents,
            "conversation_id": conversation_id,
            "new_conversation": new_conversation,
            "turn_id": turn_id,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Widget chat error: {e}")
        raise HTTPException(status_code=500, detail="Chat service unavailable")


@app.get("/list-documents")
async def list_documents():
    """List all documents in the Docs bucket"""
    try:
        response = supabase.storage.from_("Docs").list("")
        logger.info(f"📋 Listed {len(response)} documents")
        return {"documents": response}
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/url-activities")
async def list_url_activities(limit: int = 20, authorization: str = Header(None)):
    """Fetch recent URL ingestion activity logs"""
    try:
        # Extract user id from Authorization header and only return that user's activities
        user_id = _extract_user_id_from_auth(authorization)

        # Diagnostic logging to help debug why users may see no activities
        logger.info("/url-activities requested by user_id=%s", user_id)

        # Fetch recent activities and filter in Python to avoid DB JSON filtering quirks
        response = (
            supabase.table(URL_ACTIVITY_TABLE)
            .select("*")
            .order("started_at", desc=True)
            .limit(limit)
            .execute()
        )

        activities = response.data if response and hasattr(response, "data") else []
        logger.info("/url-activities fetched %d activities (before filtering)", len(activities))
        try:
            logger.info("/url-activities sample metadata: %s", [a.get("metadata") for a in activities[:3]])
        except Exception:
            pass

        def _belongs_to_user(activity: dict, uid: str) -> bool:
            # Check top-level fields first
            if not activity:
                return False
            if activity.get("user_id") == uid or activity.get("site_id") == uid:
                return True

            # Check metadata (may be None or json)
            meta = activity.get("metadata") or {}
            try:
                if isinstance(meta, str):
                    import json

                    meta = json.loads(meta)
            except Exception:
                meta = {}

            if isinstance(meta, dict):
                if str(meta.get("user_id")) == str(uid) or str(meta.get("site_id")) == str(uid):
                    return True

            return False

        filtered = [a for a in activities if _belongs_to_user(a, user_id)]
        return {"activities": filtered}
    except Exception as e:
        logger.error(f"Error fetching URL activities: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/widget.js")
async def widget_script():
    script = f"""
    (function() {{
    const config = window.supportBotConfig || {{}};
    const WIDGET_ID = config.widgetId || 'default';
      const API_BASE = config.apiBase || '{WIDGET_SCRIPT_BASE_URL}';
    const WIDGET_CONTAINER_ID = 'supportbot-widget-container';

            if (document.getElementById(WIDGET_CONTAINER_ID)) {{
                return;
            }}

      const styles = document.createElement('style');
      styles.textContent = `
        .supportbot-widget-bubble {{
          position: fixed;
          bottom: 24px;
          z-index: 9998;
          width: 56px;
          height: 56px;
          border-radius: 50%;
          display: flex;
          align-items: center;
          justify-content: center;
          box-shadow: 0 10px 40px rgba(15,23,42,0.18);
          cursor: pointer;
          transition: transform 0.2s ease;
        }}

        .supportbot-widget-bubble:hover {{
          transform: translateY(-3px);
        }}

        .supportbot-chat-window {{
          position: fixed;
          bottom: 90px;
          width: 360px;
          height: 520px;
          border-radius: 16px;
          box-shadow: 0 20px 60px rgba(15,23,42,0.25);
          overflow: hidden;
          display: flex;
          flex-direction: column;
          background: #ffffff;
          z-index: 9999;
        }}
      `;
      document.head.appendChild(styles);

    const container = document.createElement('div');
    container.id = WIDGET_CONTAINER_ID;
      document.body.appendChild(container);

      const state = {{
        open: false,
        messages: [],
        loading: false,
        conversationId: null,
        feedbackByTurn: {{}},
      }};

      function render() {{
        const primaryColor = config.primaryColor || '#3B82F6';
        const backgroundColor = config.backgroundColor || '#FFFFFF';
        const position = config.position === 'bottom-left' ? 'left' : 'right';
        const welcomeMessage = config.welcomeMessage || 'Hi there! How can I help you?';
        const placeholder = config.placeholder || 'Type your message...';
        const showBranding = config.showBranding !== false;

        const bubble = document.createElement('div');
        bubble.className = 'supportbot-widget-bubble';
        bubble.style.background = primaryColor;
        bubble.style[position] = '24px';
        bubble.textContent = '💬';
        bubble.onclick = () => toggle(true);

        const windowEl = document.createElement('div');
        windowEl.className = 'supportbot-chat-window';
        windowEl.style[position] = '24px';
        windowEl.style.display = state.open ? 'flex' : 'none';

        windowEl.innerHTML = `
          <div style="background:${{config.primaryColor || '#3B82F6'}};color:${{config.backgroundColor || '#FFFFFF'}};padding:16px;">
            <div style="display:flex;justify-content:space-between;align-items:center;">
              <strong>Support</strong>
              <button aria-label="Close chat" style="background:none;border:none;color:${{config.backgroundColor || '#FFFFFF'}};font-size:20px;cursor:pointer;">×</button>
            </div>
            <p style="margin-top:8px;font-size:13px;opacity:0.8;">We're here to help</p>
          </div>
          <div class="supportbot-messages" style="flex:1;padding:16px;overflow-y:auto;background:${{config.backgroundColor || '#FFFFFF'}};"></div>
          <div class="supportbot-input" style="padding:16px;border-top:1px solid #e2e8f0;background:${{config.backgroundColor || '#FFFFFF'}};">
            <form style="display:flex;gap:8px;">
              <input type="text" placeholder="${{config.placeholder || 'Type your message...'}}" style="flex:1;padding:10px 12px;border:1px solid #cbd5f5;border-radius:8px;" />
              <button type="submit" style="background:${{config.primaryColor || '#3B82F6'}};color:${{config.backgroundColor || '#FFFFFF'}};border:none;border-radius:8px;padding:0 14px;cursor:pointer;">Send</button>
            </form>
          </div>
        `;
        if (showBranding) {{
          const inputSection = windowEl.querySelector('.supportbot-input');
          if (inputSection) {{
            inputSection.insertAdjacentHTML(
              'beforeend',
              '<p style="margin-top:8px;text-align:center;font-size:11px;color:#94a3b8;">Powered by SupportBot</p>'
            );
          }}
        }}

        const closeButton = windowEl.querySelector('button');
        if (closeButton) {{
          closeButton.addEventListener('click', () => toggle(false));
        }}

        const form = windowEl.querySelector('form');
        const input = windowEl.querySelector('input');
        const messagesContainer = windowEl.querySelector('.supportbot-messages');

        if (state.messages.length === 0) {{
          state.messages.push({{ id: 'welcome', role: 'assistant', content: welcomeMessage }});
        }}

        if (messagesContainer) {{
          messagesContainer.innerHTML = '';
          state.messages.forEach((msg) => {{
            const bubble = document.createElement('div');
            bubble.style.marginBottom = '12px';
            bubble.style.display = 'flex';
            bubble.style.justifyContent = msg.role === 'assistant' ? 'flex-start' : 'flex-end';

            const bubbleInner = document.createElement('div');
            bubbleInner.style.maxWidth = '80%';
            bubbleInner.style.padding = '12px';
            bubbleInner.style.borderRadius = '14px';
            bubbleInner.style.background = msg.role === 'assistant' ? primaryColor + '15' : primaryColor;
            bubbleInner.style.color = msg.role === 'assistant' ? '#0f172a' : backgroundColor;
            bubbleInner.style.fontSize = '14px';
            bubbleInner.textContent = msg.content;

            if (msg.role === 'assistant' && msg.turnId) {{
              const feedbackRow = document.createElement('div');
              feedbackRow.style.marginTop = '8px';
              feedbackRow.style.display = 'flex';
              feedbackRow.style.gap = '8px';

              const makeButton = (label, sentiment, activeColor) => {{
                const button = document.createElement('button');
                button.type = 'button';
                button.style.display = 'inline-flex';
                button.style.alignItems = 'center';
                button.style.gap = '4px';
                button.style.fontSize = '12px';
                button.style.border = 'none';
                button.style.background = 'transparent';
                button.style.cursor = 'pointer';
                button.style.color = state.feedbackByTurn[msg.turnId] === sentiment ? activeColor : '#64748b';
                button.textContent = label;
                button.onclick = () => submitFeedback(msg.turnId, sentiment);
                return button;
              }};

              feedbackRow.appendChild(makeButton('👍 Helpful', 'positive', primaryColor));
              feedbackRow.appendChild(makeButton('👎 Not helpful', 'negative', '#ef4444'));
              bubbleInner.appendChild(feedbackRow);
            }}

            bubble.appendChild(bubbleInner);
            messagesContainer.appendChild(bubble);
          }});

          if (state.loading) {{
            const thinking = document.createElement('div');
            thinking.style.fontSize = '12px';
            thinking.style.color = '#64748b';
            thinking.textContent = 'Thinking…';
            messagesContainer.appendChild(thinking);
          }}

          messagesContainer.scrollTop = messagesContainer.scrollHeight;
        }}

        if (form && input) {{
          form.onsubmit = async (event) => {{
            event.preventDefault();
            const userMessage = input.value.trim();
            if (!userMessage || state.loading) {{
              return;
            }}

            state.messages.push({{ id: 'user-' + Date.now(), role: 'user', content: userMessage }});
            input.value = '';
            state.loading = true;
            rerender();

            try {{
                            const tokenResponse = await fetch(API_BASE + '/widget/token', {{
                                method: 'POST',
                                headers: {{ 'Content-Type': 'application/json' }},
                                body: JSON.stringify({{ widget_id: WIDGET_ID }}),
                            }});

              if (!tokenResponse.ok) {{
                throw new Error('Token request failed');
              }}

              const {{ token }} = await tokenResponse.json();

              const chatResponse = await fetch(API_BASE + '/widget/chat', {{
                method: 'POST',
                headers: {{
                  'Content-Type': 'application/json',
                  'Authorization': 'Bearer ' + token,
                }},
                body: JSON.stringify({{
                  query: userMessage,
                  history: state.messages,
                  top_k: config.topK || 5,
                  temperature: config.temperature || 0.2,
                  conversation_id: state.conversationId,
                }}),
              }});

              if (!chatResponse.ok) {{
                throw new Error('Chat request failed');
              }}

              const data = await chatResponse.json();

              if (data.conversation_id) {{
                state.conversationId = data.conversation_id;
              }}

              const turnId = data.turn_id || 'assistant-' + Date.now();
              state.messages.push({{
                id: turnId,
                role: 'assistant',
                content: data.answer || "I'm sorry, I could not come up with an answer.",
                turnId,
              }});
            }} catch (error) {{
              console.error('Widget error:', error);
              state.messages.push({{
                id: 'error-' + Date.now(),
                role: 'assistant',
                content: 'Something went wrong while fetching the answer. Please try again later.',
              }});
            }} finally {{
              state.loading = false;
              rerender();
            }}
          }};
        }}

        container.innerHTML = '';
        container.appendChild(bubble);
        container.appendChild(windowEl);
      }}

      function rerender() {{
        render();
      }}

      function toggle(open) {{
        state.open = open;
        rerender();
      }}

      async function submitFeedback(turnId, sentiment) {{
        state.feedbackByTurn[turnId] = sentiment;
        rerender();

        try {{
                            const tokenResponse = await fetch(API_BASE + '/widget/token', {{
                                method: 'POST',
                                headers: {{ 'Content-Type': 'application/json' }},
                                body: JSON.stringify({{ widget_id: WIDGET_ID }}),
                            }});

          if (!tokenResponse.ok) {{
            throw new Error('Feedback token request failed');
          }}

          const {{ token }} = await tokenResponse.json();

          await fetch(API_BASE + '/analytics/feedback', {{
            method: 'POST',
            headers: {{
              'Content-Type': 'application/json',
              'Authorization': 'Bearer ' + token,
            }},
                        body: JSON.stringify({{
                            widget_id: WIDGET_ID,
                            conversation_id: state.conversationId,
                            turn_id: turnId,
                            sentiment,
                            metadata: {{ source: 'widget' }},
                        }}),
          }});
        }} catch (error) {{
          console.error('Feedback submission failed:', error);
        }}
      }}

      render();
    }})(); """

    return Response(content=script, media_type="application/javascript")


@app.get("/debug-file/{file_path:path}")
async def debug_file(file_path: str):
    """Debug endpoint to test file download"""
    try:
        logger.info(f"🔍 Debugging file: {file_path}")

        # List all files first
        all_files = supabase.storage.from_("Docs").list("")
        logger.info(f"Available files: {[f['name'] for f in all_files]}")

        # Try to download
        try:
            response = supabase.storage.from_("Docs").download(file_path)
            if response:
                logger.info(f"✅ Successfully downloaded: {file_path}")
                return {
                    "status": "success",
                    "message": f"File {file_path} downloaded successfully",
                    "size": len(response),
                }
            else:
                logger.error(f"❌ Download returned None for: {file_path}")
                return {"status": "error", "message": "Download returned None"}
        except Exception as e:
            logger.error(f"❌ Download failed: {e}")
            return {
                "status": "error",
                "message": str(e),
                "available_files": [f["name"] for f in all_files],
            }

    except Exception as e:
        logger.error(f"Debug error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# Webhook handler for Supabase storage events
@app.post("/webhook/storage")
async def storage_webhook(request: dict):
    """Handle Supabase storage webhook events"""
    try:
        logger.info(f"Webhook received: {request}")
        event_type = request.get("type")

        if event_type == "INSERT":
            # New file uploaded
            record = request.get("record", {})
            file_path = record.get("name")
            file_name = record.get("name", "").split("/")[-1]

            if file_path and file_name:
                logger.info(f"🚀 New file uploaded: {file_name}")
                logger.info(f"📁 File path: {file_path}")
                # Skip processing if we've already processed this document
                try:
                    already = await pipeline.is_document_processed(file_path)
                except Exception:
                    already = False

                if already:
                    logger.info(f"⏭️ Skipping webhook processing for already-processed file: {file_path}")
                else:
                    # Process document in background
                    asyncio.create_task(pipeline.process_document_by_path(file_path))
            else:
                logger.warning(f"Invalid webhook data: {request}")
        else:
            logger.info(f"Webhook event type: {event_type} (not processing)")

        return {"status": "success"}

    except Exception as e:
        logger.error(f"Webhook error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
